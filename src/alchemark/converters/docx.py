"""DOCX → Markdown converter."""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING, Any, ClassVar

from alchemark.converters.base import BaseConverter
from alchemark.core import Result
from alchemark.exceptions import ConversionError

if TYPE_CHECKING:
    pass


# OOXML namespaces used when scanning runs for inline images.
_NS_DRAWING = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
_NS_RELS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"


class DocxConverter(BaseConverter):
    """Converts Microsoft Word (.docx) documents to Markdown."""

    extensions: ClassVar[list[str]] = [".docx"]
    name: ClassVar[str] = "docx"

    def convert(self, path: Path) -> Result:
        try:
            from docx import Document
            from docx.document import Document as DocumentType
        except ImportError as e:  # pragma: no cover
            raise ConversionError(
                path,
                "python-docx is not installed",
                hint="Install it with: pip install python-docx",
            ) from e

        try:
            doc: DocumentType = Document(str(path))
        except Exception as e:
            raise ConversionError(
                path,
                f"Could not open document: {e}",
                hint="The file may be corrupted or password-protected.",
            ) from e

        warnings: list[str] = []
        lines: list[str] = []

        # Optionally extract embedded images and build an rId → relative-path map.
        rid_to_path: dict[str, str] = {}
        if self.preserve_images:
            try:
                rid_to_path = self._extract_images(doc, path)
            except Exception as e:
                warnings.append(f"Image extraction failed: {e}. Continuing without images.")
                rid_to_path = {}

        for para in doc.paragraphs:
            style = (para.style.name if para.style else "").lower()

            # Build paragraph text, possibly with inline image markers.
            body = self._format_runs(para, rid_to_path)
            text = body.strip()

            if not text:
                lines.append("")
                continue

            if style.startswith("heading"):
                # "heading 1" -> 1, "heading 2" -> 2, etc.
                try:
                    level = int(style.split()[-1])
                    level = max(1, min(level, 6))
                except (ValueError, IndexError):
                    level = 1
                lines.append(f"{'#' * level} {text}")
            elif "list" in style or "bullet" in style:
                lines.append(f"- {text}")
            else:
                lines.append(body)

        # Tables
        for table in doc.tables:
            lines.append("")
            lines.extend(self._table_to_markdown(table))
            lines.append("")

        # Metadata
        core_props = doc.core_properties
        metadata: dict[str, Any] = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "word_count": sum(len(p.text.split()) for p in doc.paragraphs),
            "image_count": len(rid_to_path),
        }

        markdown = "\n".join(lines).strip() + "\n"
        # Collapse 3+ blank lines into 2
        while "\n\n\n" in markdown:
            markdown = markdown.replace("\n\n\n", "\n\n")

        return Result(
            markdown=markdown,
            source=path,
            metadata=metadata,
            warnings=warnings,
        )

    def _extract_images(self, doc: Any, source: Path) -> dict[str, str]:
        """Extract every embedded image to disk, return mapping of rId → relative path.

        Images are written into ``self.image_dir`` if set, otherwise into
        ``<source_stem>_images/`` next to the source file. Markdown references
        use a path relative to the source's parent directory so the result is
        portable.
        """
        from docx.parts.image import ImagePart

        if self.image_dir is not None:
            target_dir = Path(self.image_dir)
        else:
            target_dir = source.parent / f"{source.stem}_images"
        target_dir.mkdir(parents=True, exist_ok=True)

        rid_to_path: dict[str, str] = {}
        counter = 0
        for rid, part in doc.part.related_parts.items():
            if not isinstance(part, ImagePart):
                continue
            counter += 1
            ext = Path(part.partname).suffix or ".bin"
            filename = f"image_{counter:03d}{ext}"
            out = target_dir / filename
            out.write_bytes(part.blob)
            # Build a markdown-friendly relative path from the source's parent.
            try:
                rel = out.relative_to(source.parent)
                rid_to_path[rid] = rel.as_posix()
            except ValueError:
                # image_dir was outside source.parent — fall back to absolute.
                rid_to_path[rid] = out.as_posix()
        return rid_to_path

    @staticmethod
    def _format_runs(para: Any, rid_to_path: dict[str, str] | None = None) -> str:
        """Format a paragraph with inline bold/italic markers and image refs.

        If ``rid_to_path`` is provided, any inline drawings whose blip rId is
        in the map are rendered as ``![](path)`` at their position.
        """
        rid_to_path = rid_to_path or {}
        parts: list[str] = []
        for run in para.runs:
            # 1. Inline images come first, in order they appear within the run.
            for blip in run._element.iter(f"{_NS_DRAWING}blip"):
                rid = blip.get(f"{_NS_RELS}embed")
                if rid and rid in rid_to_path:
                    parts.append(f"[![]({rid_to_path[rid]})]({rid_to_path[rid]})")
            # 2. Then the run's text with formatting.
            text = run.text
            if not text:
                continue
            if run.bold and run.italic:
                parts.append(f"***{text}***")
            elif run.bold:
                parts.append(f"**{text}**")
            elif run.italic:
                parts.append(f"*{text}*")
            else:
                parts.append(text)
        if parts:
            return "".join(parts)
        return str(para.text)

    @staticmethod
    def _table_to_markdown(table: Any) -> list[str]:
        """Convert a docx Table to Markdown table rows."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") or " " for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if not rows:
            return []

        # Insert separator after the first (header) row
        col_count = rows[0].count("|") - 1
        separator = "| " + " | ".join(["---"] * col_count) + " |"
        return [rows[0], separator, *rows[1:]]
