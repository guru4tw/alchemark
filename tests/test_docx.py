"""Tests for the DOCX converter."""

from __future__ import annotations

from pathlib import Path

import pytest

from alchemark import alchemize
from alchemark.core import Alchemist


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a small sample .docx file for testing."""
    pytest.importorskip("docx")
    from docx import Document

    doc = Document()
    doc.core_properties.title = "Test Doc"
    doc.core_properties.author = "Tester"

    doc.add_heading("Main Title", level=1)
    doc.add_heading("Subsection", level=2)
    doc.add_paragraph("This is a normal paragraph.")
    p = doc.add_paragraph()
    p.add_run("Bold text").bold = True
    p.add_run(" and ")
    p.add_run("italic").italic = True

    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "2"

    out = tmp_path / "sample.docx"
    doc.save(str(out))
    return out


def test_docx_headings(sample_docx: Path) -> None:
    md = alchemize(sample_docx)
    assert "# Main Title" in md
    assert "## Subsection" in md


def test_docx_paragraph(sample_docx: Path) -> None:
    md = alchemize(sample_docx)
    assert "This is a normal paragraph." in md


def test_docx_inline_formatting(sample_docx: Path) -> None:
    md = alchemize(sample_docx)
    assert "**Bold text**" in md
    assert "*italic*" in md


def test_docx_tables(sample_docx: Path) -> None:
    md = alchemize(sample_docx)
    assert "| A | B |" in md
    assert "| 1 | 2 |" in md
    assert "| --- | --- |" in md


def test_docx_metadata(sample_docx: Path) -> None:
    a = Alchemist()
    result = a.transmute(sample_docx)
    assert result.metadata["title"] == "Test Doc"
    assert result.metadata["author"] == "Tester"
