"""Tests for DOCX image extraction (Roadmap: preserve_images for DOCX)."""

from __future__ import annotations

import io
from pathlib import Path

import pytest

from alchemark import Alchemist

pytest.importorskip("docx")
pytest.importorskip("PIL")


def _make_png_bytes(color: tuple[int, int, int] = (0, 128, 255)) -> bytes:
    """Tiny in-memory PNG for embedding."""
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (40, 20), color=color).save(buf, format="PNG")
    return buf.getvalue()


@pytest.fixture
def docx_with_image(tmp_path: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_heading("Doc With Image", level=1)
    doc.add_paragraph("Some text before the picture.")

    img_bytes = _make_png_bytes()
    img_path = tmp_path / "_seed.png"
    img_path.write_bytes(img_bytes)
    doc.add_picture(str(img_path))

    doc.add_paragraph("Some text after the picture.")
    out = tmp_path / "with_image.docx"
    doc.save(str(out))
    return out


def test_docx_preserve_images_default_dir(docx_with_image: Path) -> None:
    a = Alchemist(preserve_images=True)
    result = a.transmute(docx_with_image)

    # Default dir: <source_stem>_images/ next to the source.
    expected_dir = docx_with_image.parent / f"{docx_with_image.stem}_images"
    assert expected_dir.is_dir()
    images = list(expected_dir.iterdir())
    assert len(images) == 1
    assert images[0].suffix.lower() in (".png", ".jpg", ".jpeg")
    # Markdown should reference the image relatively.
    assert f"{expected_dir.name}/" in result.markdown
    assert "![](" in result.markdown
    # Metadata should record the image count.
    assert result.metadata["image_count"] == 1


def test_docx_preserve_images_custom_dir(tmp_path: Path, docx_with_image: Path) -> None:
    target = tmp_path / "custom_imgs"
    a = Alchemist(preserve_images=True, image_dir=target)
    result = a.transmute(docx_with_image)

    assert target.is_dir()
    images = list(target.iterdir())
    assert len(images) == 1
    # Reference appears in markdown — exact prefix depends on relative location.
    assert "![](" in result.markdown
    assert result.metadata["image_count"] == 1


def test_docx_preserve_images_false_does_not_extract(tmp_path: Path, docx_with_image: Path) -> None:
    a = Alchemist(preserve_images=False)
    result = a.transmute(docx_with_image)

    # No image dir created.
    assert not (docx_with_image.parent / f"{docx_with_image.stem}_images").exists()
    # No image references in markdown.
    assert "![](" not in result.markdown
    assert result.metadata["image_count"] == 0


def test_docx_no_warning_when_preserve_images_succeeds(docx_with_image: Path) -> None:
    a = Alchemist(preserve_images=True)
    result = a.transmute(docx_with_image)
    # Old placeholder warning should be gone.
    assert not any("not yet implemented" in w.lower() for w in result.warnings)
